"""
app/utils/file_extractor.py

Модуль для извлечения текста из различных типов файлов
Поддерживает: PDF, Word (DOC/DOCX), Excel (XLS/XLSX), TXT, RTF, CSV

Использование в main.py:
    from app.utils.file_extractor import extract_text_from_file, cleanup_file
"""

import os
import logging
from typing import Optional, Dict, Any
from pathlib import Path

# Библиотеки для работы с файлами
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

# Настройка логирования
logger = logging.getLogger(__name__)


class FileTextExtractor:
    """
    Класс для извлечения текста из различных типов файлов
    """

    def __init__(self):
        """Инициализация экстрактора с проверкой доступности библиотек"""
        self.supported_formats = {
            'pdf': PdfReader is not None,
            'docx': docx is not None,
            'doc': docx is not None,
            'xlsx': pd is not None,
            'xls': pd is not None,
            'csv': pd is not None,
            'txt': True,
            'rtf': rtf_to_text is not None
        }

        # Логируем недоступные форматы
        unavailable = [fmt for fmt, available in self.supported_formats.items() if not available]
        if unavailable:
            logger.warning(f"Следующие форматы недоступны (отсутствуют библиотеки): {', '.join(unavailable)}")

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Главная функция для извлечения текста из файла

        Args:
            file_path: Путь к файлу

        Returns:
            Dict с ключами:
                - success: bool - успешность операции
                - text: str - извлеченный текст
                - error: str - сообщение об ошибке (если есть)
                - metadata: dict - дополнительная информация о файле
        """
        try:
            # Проверка существования файла
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'text': '',
                    'error': f'Файл не найден: {file_path}',
                    'metadata': {}
                }

            # Определение расширения файла
            file_extension = Path(file_path).suffix.lower().lstrip('.')

            # Проверка поддержки формата
            if file_extension not in self.supported_formats:
                return {
                    'success': False,
                    'text': '',
                    'error': f'Неподдерживаемый формат файла: {file_extension}',
                    'metadata': {'format': file_extension}
                }

            if not self.supported_formats[file_extension]:
                return {
                    'success': False,
                    'text': '',
                    'error': f'Библиотека для обработки {file_extension} не установлена',
                    'metadata': {'format': file_extension}
                }

            # Выбор метода извлечения в зависимости от типа файла
            extraction_methods = {
                'pdf': self._extract_from_pdf,
                'docx': self._extract_from_docx,
                'doc': self._extract_from_docx,
                'xlsx': self._extract_from_excel,
                'xls': self._extract_from_excel,
                'csv': self._extract_from_csv,
                'txt': self._extract_from_txt,
                'rtf': self._extract_from_rtf
            }

            method = extraction_methods.get(file_extension)
            if method:
                result = method(file_path)
                result['metadata']['format'] = file_extension
                return result

            return {
                'success': False,
                'text': '',
                'error': f'Метод извлечения для {file_extension} не реализован',
                'metadata': {'format': file_extension}
            }

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из {file_path}: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Непредвиденная ошибка: {str(e)}',
                'metadata': {}
            }

    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из PDF файла"""
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1}: {str(e)}")

            full_text = '\n\n'.join(text_parts)

            return {
                'success': True,
                'text': full_text.strip(),
                'error': None,
                'metadata': {
                    'pages': len(reader.pages),
                    'extracted_pages': len(text_parts)
                }
            }

        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Ошибка обработки PDF: {str(e)}',
                'metadata': {}
            }

    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из Word документа (DOCX/DOC)"""
        try:
            doc = docx.Document(file_path)
            text_parts = []

            # Извлечение текста из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = '\n'.join(text_parts)

            return {
                'success': True,
                'text': full_text.strip(),
                'error': None,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables)
                }
            }

        except Exception as e:
            logger.error(f"Ошибка обработки Word: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Ошибка обработки Word документа: {str(e)}',
                'metadata': {}
            }

    def _extract_from_excel(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из Excel файла (XLSX/XLS)"""
        try:
            # Чтение всех листов
            excel_file = pd.ExcelFile(file_path)
            text_parts = []

            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)

                    # Добавляем название листа
                    text_parts.append(f"=== Лист: {sheet_name} ===")

                    # Конвертируем DataFrame в текст
                    sheet_text = df.to_string(index=False, na_rep='')
                    text_parts.append(sheet_text)
                    text_parts.append('')  # Пустая строка между листами

                except Exception as e:
                    logger.warning(f"Ошибка чтения листа {sheet_name}: {str(e)}")

            full_text = '\n'.join(text_parts)

            return {
                'success': True,
                'text': full_text.strip(),
                'error': None,
                'metadata': {
                    'sheets': len(excel_file.sheet_names),
                    'sheet_names': excel_file.sheet_names
                }
            }

        except Exception as e:
            logger.error(f"Ошибка обработки Excel: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Ошибка обработки Excel файла: {str(e)}',
                'metadata': {}
            }

    def _extract_from_csv(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из CSV файла"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'latin1']
            df = None
            used_encoding = None

            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                return {
                    'success': False,
                    'text': '',
                    'error': 'Не удалось определить кодировку CSV файла',
                    'metadata': {}
                }

            # Конвертируем DataFrame в текст
            text = df.to_string(index=False, na_rep='')

            return {
                'success': True,
                'text': text.strip(),
                'error': None,
                'metadata': {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'encoding': used_encoding,
                    'column_names': list(df.columns)
                }
            }

        except Exception as e:
            logger.error(f"Ошибка обработки CSV: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Ошибка обработки CSV файла: {str(e)}',
                'metadata': {}
            }

    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из TXT файла"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'latin1']
            text = None
            used_encoding = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                return {
                    'success': False,
                    'text': '',
                    'error': 'Не удалось определить кодировку TXT файла',
                    'metadata': {}
                }

            return {
                'success': True,
                'text': text.strip(),
                'error': None,
                'metadata': {
                    'encoding': used_encoding,
                    'lines': len(text.splitlines())
                }
            }

        except Exception as e:
            logger.error(f"Ошибка обработки TXT: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Ошибка обработки TXT файла: {str(e)}',
                'metadata': {}
            }

    def _extract_from_rtf(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из RTF файла"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()

            # Конвертируем RTF в обычный текст
            text = rtf_to_text(rtf_content)

            return {
                'success': True,
                'text': text.strip(),
                'error': None,
                'metadata': {
                    'original_size': len(rtf_content),
                    'extracted_size': len(text)
                }
            }

        except Exception as e:
            logger.error(f"Ошибка обработки RTF: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': f'Ошибка обработки RTF файла: {str(e)}',
                'metadata': {}
            }


# Создание глобального экземпляра экстрактора
file_extractor = FileTextExtractor()


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Удобная функция для извлечения текста из файла

    Args:
        file_path: Путь к файлу

    Returns:
        Dict с результатом извлечения
    """
    return file_extractor.extract_text(file_path)


def cleanup_file(file_path: str) -> bool:
    """
    Удаление файла после обработки

    Args:
        file_path: Путь к файлу

    Returns:
        bool: True если файл успешно удален, False в противном случае
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Файл успешно удален: {file_path}")
            return True
        else:
            logger.warning(f"Файл не найден для удаления: {file_path}")
            return False
    except Exception as e:
        logger.error(f"Ошибка при удалении файла {file_path}: {str(e)}")
        return False