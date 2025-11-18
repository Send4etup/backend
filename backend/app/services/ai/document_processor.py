# backend/services/ai/document_processor.py
"""
Модуль для обработки документов различных форматов
Включает извлечение текста из PDF, Word, Excel, CSV и других форматов
"""

import logging
from pathlib import Path
from typing import Optional, Union
import PyPDF2
import docx
import pandas as pd
from openpyxl import load_workbook
import asyncio

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Класс для обработки документов различных форматов"""

    def __init__(self, max_text_length: int = 10e10):
        """
        Инициализация процессора документов

        Args:
            max_text_length: Максимальная длина извлекаемого текста
        """
        self.max_text_length = max_text_length

        # Поддерживаемые форматы документов
        self.supported_formats = {
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'excel': ['.xlsx', '.xls'],
            'csv': ['.csv'],
            'text': ['.txt', '.md', '.json', '.xml', '.log']
        }

    async def extract_text_from_pdf(
            self,
            file_path: str,
            max_pages: Optional[int] = None
    ) -> str:
        """
        Извлечение текста из PDF файла

        Args:
            file_path: Путь к PDF файлу
            max_pages: Максимальное количество страниц для обработки

        Returns:
            Извлеченный текст
        """
        try:
            text = ""
            file_name = Path(file_path).name

            logger.info(f"Extracting text from PDF: {file_name}")

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

                # Определяем количество страниц для обработки
                pages_to_process = min(
                    total_pages,
                    max_pages if max_pages else total_pages
                )

                logger.info(
                    f"PDF has {total_pages} pages, "
                    f"processing {pages_to_process} pages"
                )

                # Извлекаем текст со страниц
                for page_num in range(pages_to_process):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()

                        if page_text:
                            text += f"\n--- Страница {page_num + 1} ---\n"
                            text += page_text + "\n"

                        # Прерываем если достигли лимита текста
                        if len(text) > self.max_text_length:
                            logger.info(
                                f"Reached text length limit at page {page_num + 1}"
                            )
                            break

                    except Exception as page_error:
                        logger.warning(
                            f"Error extracting text from page {page_num + 1}: "
                            f"{page_error}"
                        )
                        continue

            # Обрезаем текст если он слишком длинный
            if len(text) > self.max_text_length:
                text = text[:self.max_text_length]
                text += "\n\n... [текст обрезан по лимиту длины]"

            if not text.strip():
                return f"PDF файл '{file_name}' не содержит извлекаемого текста или текст зашифрован."

            logger.info(
                f"Successfully extracted {len(text)} characters from PDF {file_name}"
            )

            return text

        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {e}")
            return f"Ошибка при чтении PDF файла: {str(e)}"

    async def extract_text_from_docx(
            self,
            file_path: str,
            max_paragraphs: Optional[int] = 100
    ) -> str:
        """
        Извлечение текста из Word документа

        Args:
            file_path: Путь к DOCX файлу
            max_paragraphs: Максимальное количество параграфов

        Returns:
            Извлеченный текст
        """
        try:
            file_name = Path(file_path).name

            logger.info(f"Extracting text from DOCX: {file_name}")

            doc = docx.Document(file_path)
            text = ""

            # Извлекаем параграфы
            total_paragraphs = len(doc.paragraphs)
            paragraphs_to_process = min(
                total_paragraphs,
                max_paragraphs if max_paragraphs else total_paragraphs
            )

            logger.info(
                f"DOCX has {total_paragraphs} paragraphs, "
                f"processing {paragraphs_to_process}"
            )

            for i, paragraph in enumerate(doc.paragraphs[:paragraphs_to_process]):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

                # Прерываем если достигли лимита
                if len(text) > self.max_text_length:
                    logger.info(
                        f"Reached text length limit at paragraph {i + 1}"
                    )
                    break

            # Извлекаем текст из таблиц
            if doc.tables:
                text += "\n--- Таблицы в документе ---\n"

                for table_idx, table in enumerate(doc.tables[:5]):  # Первые 5 таблиц
                    text += f"\nТаблица {table_idx + 1}:\n"

                    try:
                        for row in table.rows[:10]:  # Первые 10 строк
                            row_text = " | ".join(
                                cell.text.strip() for cell in row.cells
                            )
                            if row_text:
                                text += row_text + "\n"
                    except Exception as table_error:
                        logger.warning(
                            f"Error extracting table {table_idx + 1}: {table_error}"
                        )

                    if len(text) > self.max_text_length:
                        break

            # Обрезаем текст если он слишком длинный
            if len(text) > self.max_text_length:
                text = text[:self.max_text_length]
                text += "\n\n... [текст обрезан по лимиту длины]"

            if not text.strip():
                return f"Word документ '{file_name}' не содержит текста."

            logger.info(
                f"Successfully extracted {len(text)} characters from DOCX {file_name}"
            )

            return text

        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {e}")
            return f"Ошибка при чтении Word документа: {str(e)}"

    async def extract_text_from_excel(
            self,
            file_path: str,
            max_rows_per_sheet: Optional[int] = 10000
    ) -> str:
        """
        Асинхронное извлечение текста из Excel с аккуратным форматированием.
        """
        try:
            file_name = Path(file_path).name
            logger.info(f"Reading Excel: {file_name}")

            wb = await asyncio.to_thread(load_workbook, file_path, data_only=True)
            text_parts = []

            for sheet in wb.worksheets[:5]:
                text_parts.append(f"\n📄 Лист: {sheet.title}\n{'-' * 40}")

                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if max_rows_per_sheet and i >= max_rows_per_sheet:
                        text_parts.append("... [остальные строки пропущены]")
                        break

                    values = [str(cell).strip() for cell in row if cell not in (None, "")]
                    if values:
                        text_parts.append(" | ".join(values))

                    if sum(len(p) for p in text_parts) > self.max_text_length:
                        text_parts.append("... [текст обрезан по лимиту]")
                        break

            text = "\n".join(text_parts)
            return text if text.strip() else f"Excel файл '{file_name}' не содержит текстовых данных."

        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {e}")
            return f"Ошибка при чтении Excel файла: {str(e)}"

    async def extract_text_from_csv(
            self,
            file_path: str,
            max_rows: int = 50,
            encoding: str = 'utf-8'
    ) -> str:
        """
        Извлечение данных из CSV файла

        Args:
            file_path: Путь к CSV файлу
            max_rows: Максимальное количество строк
            encoding: Кодировка файла

        Returns:
            Описание данных из CSV
        """
        try:
            file_name = Path(file_path).name

            logger.info(f"Extracting data from CSV: {file_name}")

            # Пробуем разные кодировки если utf-8 не работает
            encodings_to_try = [encoding, 'utf-8', 'latin-1', 'cp1251']
            df = None
            used_encoding = None

            for enc in encodings_to_try:
                try:
                    df = pd.read_csv(file_path, nrows=max_rows, encoding=enc)
                    used_encoding = enc
                    logger.info(f"Successfully read CSV with encoding: {enc}")
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                return f"Не удалось прочитать CSV файл '{file_name}' ни с одной из кодировок."

            description = f"CSV файл '{file_name}' (кодировка: {used_encoding})\n"
            description += f"Размер: {len(df)} строк, {len(df.columns)} столбцов\n"
            description += f"Столбцы: {', '.join(df.columns.tolist())}\n\n"

            # Добавляем информацию о типах данных
            description += "Типы данных:\n"
            for col, dtype in df.dtypes.items():
                non_null = df[col].notna().sum()
                description += f"  - {col}: {dtype} (заполнено: {non_null}/{len(df)})\n"

            # Статистика по числовым колонкам
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                description += "\nСтатистика по числовым столбцам:\n"
                stats = df[numeric_cols].describe()
                description += stats.to_string(max_cols=5) + "\n"

            # Первые строки данных
            description += f"\nПервые {min(10, len(df))} строк данных:\n"
            description += df.head(10).to_string(
                max_cols=10,
                max_colwidth=50,
                index=True
            )

            # Обрезаем если слишком длинно
            if len(description) > self.max_text_length:
                description = description[:self.max_text_length]
                description += "\n\n... [данные обрезаны по лимиту длины]"

            logger.info(
                f"Successfully extracted data from CSV {file_name}"
            )

            return description

        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {e}")
            return f"Ошибка при чтении CSV файла: {str(e)}"

    async def extract_text_from_text_file(
            self,
            file_path: str,
            encoding: str = 'utf-8'
    ) -> str:
        """
        Извлечение текста из текстового файла

        Args:
            file_path: Путь к текстовому файлу
            encoding: Кодировка файла

        Returns:
            Содержимое файла
        """
        try:
            file_name = Path(file_path).name

            logger.info(f"Reading text file: {file_name}")

            # Пробуем разные кодировки
            encodings_to_try = [encoding, 'utf-8', 'latin-1', 'cp1251']
            content = None
            used_encoding = None

            for enc in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=enc, errors='ignore') as f:
                        content = f.read(self.max_text_length)
                    used_encoding = enc
                    logger.info(f"Successfully read text file with encoding: {enc}")
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                return f"Не удалось прочитать текстовый файл '{file_name}'."

            # Обрезаем если нужно
            if len(content) > self.max_text_length:
                content = content[:self.max_text_length]
                content += "\n\n... [текст обрезан по лимиту длины]"

            result = f"Содержимое файла '{file_name}' (кодировка: {used_encoding}):\n\n{content}"

            logger.info(
                f"Successfully read {len(content)} characters from text file {file_name}"
            )

            return result

        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return f"Ошибка при чтении текстового файла: {str(e)}"

    async def extract_text_from_file(
            self,
            file_path: str,
            file_type: str
    ) -> str:
        """
        Универсальная функция извлечения текста из файлов

        Args:
            file_path: Путь к файлу
            file_type: MIME тип файла

        Returns:
            Извлеченный текст или описание содержимого
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            file_name = Path(file_path).name

            logger.info(
                f"Extracting text from file: {file_name}, "
                f"type: {file_type}, extension: {file_extension}"
            )

            # Excel файлы (.xlsx, .xls)
            if file_extension in ['.xlsx', '.xls']:
                logger.info(f"Detected Excel file by extension: {file_extension}")
                return await self.extract_text_from_excel(file_path)

            # Word документы (.docx, .doc)
            elif file_extension in ['.docx', '.doc']:
                logger.info(f"Detected Word file by extension: {file_extension}")
                return await self.extract_text_from_docx(file_path)

            # PDF файлы
            elif file_extension == '.pdf':
                logger.info(f"Detected PDF file by extension: {file_extension}")
                return await self.extract_text_from_pdf(file_path)

            # CSV файлы
            elif file_extension == '.csv':
                logger.info(f"Detected CSV file by extension: {file_extension}")
                return await self.extract_text_from_csv(file_path)

            # Текстовые файлы
            elif file_extension in ['.txt', '.md', '.json', '.xml', '.log', '.rtf']:
                logger.info(f"Detected text file by extension: {file_extension}")
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read(5000)  # Первые 5000 символов
                return content

            # Неподдерживаемый формат
            else:
                logger.warning(f"Unsupported file format: {file_type}")
                return (
                    f"Формат файла '{file_name}' ({file_type}) поддерживается для загрузки, "
                    f"но извлечение текста не реализовано."
                )

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"Ошибка при обработке файла: {str(e)}"

    def get_document_info(self, file_path: str) -> dict:
        """
        Получение информации о документе

        Args:
            file_path: Путь к файлу

        Returns:
            Словарь с информацией о документе
        """
        try:
            path = Path(file_path)
            file_size_bytes = path.stat().st_size
            file_size_mb = round(file_size_bytes / (1024 * 1024), 2)

            info = {
                'filename': path.name,
                'extension': path.suffix.lower(),
                'file_size_bytes': file_size_bytes,
                'file_size_mb': file_size_mb,
                'is_supported': self.is_supported_format(file_path)
            }

            # Дополнительная информация в зависимости от типа
            extension = path.suffix.lower()

            if extension == '.pdf':
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        info['page_count'] = len(pdf_reader.pages)
                        info['is_encrypted'] = pdf_reader.is_encrypted
                except:
                    pass

            elif extension == '.docx':
                try:
                    doc = docx.Document(file_path)
                    info['paragraph_count'] = len(doc.paragraphs)
                    info['table_count'] = len(doc.tables)
                except:
                    pass

            elif extension in ['.xlsx', '.xls']:
                try:
                    excel_file = pd.ExcelFile(file_path)
                    info['sheet_count'] = len(excel_file.sheet_names)
                    info['sheet_names'] = excel_file.sheet_names
                except:
                    pass

            elif extension == '.csv':
                try:
                    # Читаем только первую строку для получения колонок
                    df = pd.read_csv(file_path, nrows=0)
                    info['column_count'] = len(df.columns)
                    info['columns'] = df.columns.tolist()
                except:
                    pass

            logger.info(f"Document info retrieved: {info}")
            return info

        except Exception as e:
            logger.error(f"Error getting document info for {file_path}: {e}")
            return {
                'filename': Path(file_path).name,
                'error': str(e)
            }

    def is_supported_format(self, file_path: str) -> bool:
        """
        Проверка поддерживается ли формат документа

        Args:
            file_path: Путь к файлу

        Returns:
            True если формат поддерживается
        """
        extension = Path(file_path).suffix.lower()

        for format_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return True

        logger.warning(f"Unsupported document format: {extension}")
        return False

    def get_supported_formats(self) -> dict:
        """
        Получение словаря поддерживаемых форматов

        Returns:
            Словарь {тип: [расширения]}
        """
        return self.supported_formats.copy()

    def get_format_type(self, file_path: str) -> Optional[str]:
        """
        Определение типа формата документа

        Args:
            file_path: Путь к файлу

        Returns:
            Тип формата или None
        """
        extension = Path(file_path).suffix.lower()

        for format_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return format_type

        return None

    def validate_document(self, file_path: str) -> tuple[bool, Optional[str]]:
        """
        Валидация документа

        Args:
            file_path: Путь к файлу

        Returns:
            Tuple (is_valid, error_message)
        """
        try:
            # Проверка существования файла
            if not Path(file_path).exists():
                return False, "Файл не найден"

            # Проверка формата
            if not self.is_supported_format(file_path):
                return False, f"Неподдерживаемый формат: {Path(file_path).suffix}"

            # Проверка размера (не более 50 MB для документов)
            file_size_mb = Path(file_path).stat().st_size / (1024 * 1024)
            if file_size_mb > 50:
                return False, f"Файл слишком большой ({file_size_mb:.1f} MB), максимум 50 MB"

            # Базовая проверка читаемости в зависимости от типа
            extension = Path(file_path).suffix.lower()

            if extension == '.pdf':
                with open(file_path, 'rb') as f:
                    PyPDF2.PdfReader(f)

            elif extension == '.docx':
                docx.Document(file_path)

            elif extension in ['.xlsx', '.xls']:
                pd.ExcelFile(file_path)

            elif extension == '.csv':
                pd.read_csv(file_path, nrows=1)

            logger.info(f"Document validation successful: {Path(file_path).name}")
            return True, None

        except Exception as e:
            logger.error(f"Document validation error: {e}")
            return False, str(e)


# Вспомогательные функции для быстрого доступа

async def extract_document_text(
        file_path: str,
        file_type: str,
        max_length: int = 5000
) -> str:
    """
    Быстрое извлечение текста из документа

    Args:
        file_path: Путь к файлу
        file_type: MIME тип
        max_length: Максимальная длина текста

    Returns:
        Извлеченный текст
    """
    processor = DocumentProcessor(max_text_length=max_length)
    return await processor.extract_text_from_file(file_path, file_type)


def get_document_metadata(file_path: str) -> dict:
    """
    Быстрое получение метаданных документа

    Args:
        file_path: Путь к файлу

    Returns:
        Словарь с метаданными
    """
    processor = DocumentProcessor()
    return processor.get_document_info(file_path)


def check_document_valid(file_path: str) -> bool:
    """
    Быстрая проверка валидности документа

    Args:
        file_path: Путь к файлу

    Returns:
        True если валидно
    """
    processor = DocumentProcessor()
    is_valid, _ = processor.validate_document(file_path)
    return is_valid